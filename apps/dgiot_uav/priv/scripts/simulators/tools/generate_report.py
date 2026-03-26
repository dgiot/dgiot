#!/usr/bin/env python3
"""
无人机测试报告生成器 - Word模板替换版本
从Parse Server获取测试数据,使用Word模板变量替换生成报告
"""

import os
import sys
import json
import requests
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 配置
PARSE_API_URL = "http://127.0.0.1/iotapi"
NGINX_REPORTS_DIR = "/data/dgiot/nginx/html/reports"
MES_REPORT_BASE_URL = "http://172.1.2.222/reports"
TEMPLATE_PATH = "/root/gitee/dgiot/apps/dgiot_uav/priv/scripts/test_report_template.docx"

class UAVTestReportGenerator:
    """无人机测试报告生成器 - Word模板替换版本"""

    def __init__(self, session_token=None):
        """初始化报告生成器"""
        self.session_token = session_token
        self.headers = {
            'Content-Type': 'application/json',
            'sessiontoken': session_token
        } if session_token else {
            'Content-Type': 'application/json'
        }

    def get_parse_data(self, device_id):
        """从Parse Server获取设备测试数据"""
        try:
            # 获取设备基本信息
            device_url = f"{PARSE_API_URL}/classes/Device/{device_id}"
            response = requests.get(device_url, headers=self.headers)
            if response.status_code != 200:
                print(f"获取设备信息失败: {response.status_code}")
                return None

            device_data = response.json()
            content = device_data.get('content', {})

            # 获取工位信息
            station_id = content.get('station_id', None)
            station_name = content.get('station_name', '未知工位')

            # 测试结果直接从content获取
            test_result = {}
            for key in ['test_result', 'last_test_result', 'final_test_result']:
                if key in content:
                    test_result = content[key]
                    print(f"从content.{key}获取测试结果")
                    break

            if not test_result:
                if 'result' in content:
                    test_result = content['result']
                elif 'test_status' in content:
                    test_result = {'overall_result': content['test_status']}
                else:
                    test_result = {'overall_result': '待测试', 'test_items': []}

            # 通过工位查询测试项
            test_items = self.get_test_items_by_station(station_id, station_name)

            # 获取产品信息
            product_id = device_data.get('product', {}).get('objectId')
            product_data = {}
            if product_id:
                product_url = f"{PARSE_API_URL}/classes/Product/{product_id}"
                prod_response = requests.get(product_url, headers=self.headers)
                if prod_response.status_code == 200:
                    product_data = prod_response.json()
                    content.update(product_data.get('content', {}))

            # 获取遥测数据
            telemetry = {}
            try:
                devicecard_url = f"{PARSE_API_URL}/devicecard/{device_id}"
                card_response = requests.get(devicecard_url, headers=self.headers)
                if card_response.status_code == 200:
                    card_data = card_response.json()
                    telemetry = card_data.get('data', {})
                    print(f"获取到{len(telemetry)}个时序数据点")
            except Exception as e:
                print(f"获取时序数据失败: {e}")

            # 获取测试过程数据(从td子表)
            test_process_data = self.get_td_subtable_data(device_id, self.session_token)

            return {
                'device_id': device_id,
                'device_name': device_data.get('name', device_id),
                'devaddr': device_data.get('devaddr', ''),
                'product_id': product_id,
                'product_name': product_data.get('name', ''),
                'content': content,
                'station_id': station_id,
                'station_name': station_name,
                'test_result': test_result,
                'test_items': test_items,  # 通过工位查询的测试项
                'telemetry': telemetry,
                'test_process_data': test_process_data
            }

        except Exception as e:
            print(f"从Parse获取数据失败: {e}")
            import traceback
            traceback.print_exc()
            return None

    def get_test_items_by_station(self, station_id, station_name):
        """通过工位查询测试项"""
        try:
            # 构建查询: 测试项的name前缀包含工位名称
            # 例如: "总测_外观检查" 匹配工位 "总测"
            if not station_name:
                return []

            # 获取测试项产品ID
            test_item_product_id = "343cf21f82"  # 测试项产品ID

            # 查询测试项
            where_clause = json.dumps({
                "product": {
                    "__type": "Pointer",
                    "className": "Product",
                    "objectId": test_item_product_id
                },
                "name": {
                    "$regex": f"^{station_name}_"
                }
            })

            test_items_url = f"{PARSE_API_URL}/classes/Device?where={where_clause}&limit=100&order=order"
            headers = {
                'Content-Type': 'application/json'
            }
            if self.session_token:
                headers['sessiontoken'] = self.session_token

            response = requests.get(test_items_url, headers=headers)
            if response.status_code == 200:
                items_data = response.json()
                results = items_data.get('results', [])
                print(f"从工位 '{station_name}' 查询到{len(results)}个测试项")
                return results
            else:
                print(f"查询测试项失败: {response.status_code}")
                return []

        except Exception as e:
            print(f"获取测试项失败: {e}")
            return []

    def get_td_subtable_data(self, device_id, session_token):
        """从无人机的td子表获取测试过程数据"""
        try:
            where_clause = json.dumps({
                "device": {
                    "__type": "Pointer",
                    "className": "Device",
                    "objectId": device_id
                }
            })

            td_url = f"{PARSE_API_URL}/classes/td?where={where_clause}&limit=100&order=-createdAt"
            headers = {
                'Content-Type': 'application/json'
            }
            if session_token:
                headers['sessiontoken'] = session_token

            response = requests.get(td_url, headers=headers)
            if response.status_code == 200:
                td_data = response.json()
                results = td_data.get('results', [])
                print(f"从td子表获取到{len(results)}条测试过程数据")
                return results
            else:
                print(f"查询td子表失败: {response.status_code}")
                return []

        except Exception as e:
            print(f"获取td子表数据失败: {e}")
            return []

    def build_template_variables(self, report_data):
        """构建Word模板变量"""
        device_id = report_data.get('device_id', '')
        device_name = report_data.get('device_name', device_id)
        devaddr = report_data.get('devaddr', '')
        product_name = report_data.get('product_name', '')
        station_name = report_data.get('station_name', '未知工位')

        # 测试结果
        test_result = report_data.get('test_result', {})
        overall_result = test_result.get('overall_result', '待测试')

        # 统计测试项
        test_items = report_data.get('test_items', [])
        total_items = len(test_items)
        passed_items = len([i for i in test_items if i.get('result') in ['passed', '通过', 'PASSED']])
        failed_items = total_items - passed_items
        pass_rate = round(passed_items / total_items * 100, 2) if total_items > 0 else 0

        # 格式化日期时间
        now = datetime.now()
        date_str = now.strftime('%Y年%m月%d日')
        time_str = now.strftime('%H:%M:%S')
        datetime_str = now.strftime('%Y年%m月%d日 %H:%M:%S')

        # 遥测数据
        telemetry = report_data.get('telemetry', {})

        # 测试过程数据
        test_process_data = report_data.get('test_process_data', [])
        process_count = len(test_process_data)

        # 构建变量字典
        variables = {
            # 基本信息
            '{无人机编号}': device_name,
            '{设备地址}': devaddr,
            '{产品名称}': product_name,
            '{工位名称}': station_name,
            '{测试日期}': date_str,
            '{测试时间}': time_str,
            '{测试日期时间}': datetime_str,

            # 测试结果
            '{测试结果}': overall_result,
            '{测试项总数}': str(total_items),
            '{通过项数}': str(passed_items),
            '{失败项数}': str(failed_items),
            '{通过率}': f"{pass_rate}%",

            # 遥测数据
            '{电压}': f"{telemetry.get('voltage', telemetry.get('battery_voltage', 'N/A'))}V",
            '{电流}': f"{telemetry.get('current', telemetry.get('battery_current', 'N/A'))}A",
            '{温度}': f"{telemetry.get('temperature', 'N/A')}°C",
            '{气压}': f"{telemetry.get('pressure', 'N/A')}hPa",
            '{湿度}': f"{telemetry.get('humidity', 'N/A')}%",

            # 测试过程
            '{测试过程记录数}': str(process_count),
        }

        # 添加测试项表格变量
        if test_items:
            test_item_table = []
            for i, item in enumerate(test_items, 1):
                name = item.get('name', item.get('step_name', ''))
                standard = item.get('standard', item.get('criteria', '-'))
                result = item.get('result', item.get('status', '-'))

                test_item_table.append({
                    '序号': str(i),
                    '测试项目': name,
                    '测试标准': standard,
                    '测试结果': result
                })
            variables['{测试项表格}'] = test_item_table
        else:
            variables['{测试项表格}'] = []

        # 添加测试过程数据表格变量
        if test_process_data:
            process_table = []
            for i, record in enumerate(test_process_data, 1):
                created_at = record.get('createdAt', '')
                if created_at:
                    try:
                        dt = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
                        time_str = dt.strftime('%Y-%m-%d %H:%M:%S')
                    except:
                        time_str = created_at
                else:
                    time_str = '-'

                step = record.get('step', record.get('test_step', '未知步骤'))
                status = record.get('status', record.get('result', '-'))

                # 获取数据内容
                data_content = {k: v for k, v in record.items()
                               if k not in ['objectId', 'createdAt', 'updatedAt', 'device', 'ACL']}
                data_str = json.dumps(data_content, ensure_ascii=False)

                process_table.append({
                    '序号': str(i),
                    '记录时间': time_str,
                    '测试步骤': step,
                    '测试状态': status,
                    '数据内容': data_str
                })
            variables['{测试过程表格}'] = process_table
        else:
            variables['{测试过程表格}'] = []

        return variables

    def replace_template_variables(self, doc, variables):
        """替换Word模板中的变量"""
        # 替换段落中的文本
        for paragraph in doc.paragraphs:
            for placeholder, value in variables.items():
                if isinstance(value, str):
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, value)

        # 替换表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in variables.items():
                            if isinstance(value, str):
                                if placeholder in paragraph.text:
                                    paragraph.text = paragraph.text.replace(placeholder, value)

        # 替换表格 - 测试项表格
        if '{测试项表格}' in [p.text for p in doc.paragraphs]:
            self.replace_table_with_data(doc, '{测试项表格}', variables['{测试项表格}'],
                                         ['序号', '测试项目', '测试标准', '测试结果'])

        # 替换表格 - 测试过程表格
        if '{测试过程表格}' in [p.text for p in doc.paragraphs]:
            self.replace_table_with_data(doc, '{测试过程表格}', variables['{测试过程表格}'],
                                         ['序号', '记录时间', '测试步骤', '数据内容'])

        return doc

    def replace_table_with_data(self, doc, placeholder, data_list, headers):
        """用数据替换表格占位符"""
        # 找到包含占位符的段落
        placeholder_found = False
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                placeholder_found = True
                # 清空占位符文本
                paragraph.text = paragraph.text.replace(placeholder, '')

                # 找到该段落后的表格
                if len(doc.tables) > 0:
                    table = doc.tables[0]  # 假设第一个表格是要替换的

                    # 清空表格
                    for row in table.rows:
                        for cell in row.cells:
                            cell.text = ''

                    # 重新构建表格
                    if len(data_list) > 0:
                        # 设置表头
                        for i, header in enumerate(headers):
                            table.rows[0].cells[i].text = header
                            table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

                        # 填充数据
                        for i, row_data in enumerate(data_list):
                            if i + 1 < len(table.rows):
                                for j, key in enumerate(headers):
                                    if key in row_data:
                                        table.rows[i + 1].cells[j].text = str(row_data[key])
                    break

        if not placeholder_found:
            print(f"警告: 未找到占位符 '{placeholder}'")

    def generate_word_report(self, report_data):
        """使用Word模板生成报告"""
        try:
            # 检查模板文件是否存在
            if not os.path.exists(TEMPLATE_PATH):
                print(f"警告: Word模板不存在: {TEMPLATE_PATH}")
                print("使用默认报告生成方式...")
                return self.generate_default_report(report_data)

            # 加载Word模板
            print(f"加载Word模板: {TEMPLATE_PATH}")
            doc = Document(TEMPLATE_PATH)

            # 构建模板变量
            variables = self.build_template_variables(report_data)
            print(f"构建了{len(variables)}个模板变量")

            # 替换模板变量
            print("替换模板变量...")
            doc = self.replace_template_variables(doc, variables)

            # 保存Word文档
            timestamp_ms = int(datetime.now().timestamp() * 1000)
            device_id = report_data.get('device_id', 'unknown')
            word_filename = f"{timestamp_ms}.docx"

            # 创建目录
            device_dir = os.path.join(NGINX_REPORTS_DIR, device_id)
            word_dir = os.path.join(device_dir, 'word')
            os.makedirs(word_dir, exist_ok=True)

            word_filepath = os.path.join(word_dir, word_filename)
            doc.save(word_filepath)

            print(f"Word报告生成成功: {word_filepath}")

            # 生成URL
            word_url = f"{MES_REPORT_BASE_URL}/{device_id}/word/{word_filename}"

            return {
                'success': True,
                'word_filepath': word_filepath,
                'word_url': word_url,
                'word_filename': word_filename
            }

        except Exception as e:
            print(f"使用模板生成报告失败: {e}")
            import traceback
            traceback.print_exc()
            print("尝试使用默认方式生成报告...")
            return self.generate_default_report(report_data)

    def generate_default_report(self, report_data):
        """使用默认方式生成报告(模板不存在时的备用方案)"""
        try:
            doc = Document()

            # 设置中文字体
            doc.styles['Normal'].font.name = '宋体'
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 标题
            title = doc.add_heading('超近距无人机测试报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 基本信息
            doc.add_heading('一、基本信息', level=1)
            variables = self.build_template_variables(report_data)
            
            info_table = doc.add_table(rows=5, cols=2)
            info_table.style = 'Light Grid Accent 1'

            info_data = [
                ('无人机编号', variables.get('{无人机编号}', '')),
                ('工位名称', variables.get('{工位名称}', '')),
                ('测试日期', variables.get('{测试日期}', '')),
                ('测试时间', variables.get('{测试时间}', '')),
                ('测试结果', variables.get('{测试结果}', ''))
            ]

            for i, (key, value) in enumerate(info_data):
                info_table.rows[i].cells[0].text = key
                info_table.rows[i].cells[1].text = str(value)

            # 测试项
            doc.add_heading('二、测试项', level=1)
            test_items = variables.get('{测试项表格}', [])
            if test_items:
                item_table = doc.add_table(rows=len(test_items) + 1, cols=4)
                item_table.style = 'Light Grid Accent 1'
                headers = ['序号', '测试项目', '测试标准', '测试结果']
                for i, header in enumerate(headers):
                    item_table.rows[0].cells[i].text = header
                    item_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

                for i, item in enumerate(test_items):
                    item_table.rows[i + 1].cells[0].text = item.get('序号', '')
                    item_table.rows[i + 1].cells[1].text = item.get('测试项目', '')
                    item_table.rows[i + 1].cells[2].text = item.get('测试标准', '')
                    item_table.rows[i + 1].cells[3].text = item.get('测试结果', '')

            # 保存
            timestamp_ms = int(datetime.now().timestamp() * 1000)
            device_id = report_data.get('device_id', 'unknown')
            word_filename = f"{timestamp_ms}.docx"

            device_dir = os.path.join(NGINX_REPORTS_DIR, device_id)
            word_dir = os.path.join(device_dir, 'word')
            os.makedirs(word_dir, exist_ok=True)

            word_filepath = os.path.join(word_dir, word_filename)
            doc.save(word_filepath)

            word_url = f"{MES_REPORT_BASE_URL}/{device_id}/word/{word_filename}"

            return {
                'success': True,
                'word_filepath': word_filepath,
                'word_url': word_url,
                'word_filename': word_filename
            }

        except Exception as e:
            print(f"默认报告生成也失败: {e}")
            import traceback
            traceback.print_exc()
            return {'success': False, 'error': str(e)}

    def convert_to_pdf(self, word_filepath):
        """将Word转换为PDF"""
        try:
            device_dir = os.path.dirname(os.path.dirname(word_filepath))
            word_filename = os.path.basename(word_filepath)
            timestamp_ms = word_filename.replace('.docx', '')
            pdf_filename = f"{timestamp_ms}.pdf"

            pdf_dir = os.path.join(device_dir, 'pdf')
            os.makedirs(pdf_dir, exist_ok=True)

            pdf_filepath = os.path.join(pdf_dir, pdf_filename)

            # 尝试使用LibreOffice转换
            try:
                cmd = f"libreoffice --headless --convert-to pdf --outdir {pdf_dir} {word_filepath}"
                result = os.system(cmd)

                if os.path.exists(pdf_filepath):
                    print(f"PDF生成成功: {pdf_filepath}")
                    device_id = os.path.basename(device_dir)
                    pdf_url = f"{MES_REPORT_BASE_URL}/{device_id}/pdf/{pdf_filename}"
                    return {
                        'success': True,
                        'pdf_filepath': pdf_filepath,
                        'pdf_url': pdf_url,
                        'pdf_filename': pdf_filename
                    }
                else:
                    print(f"LibreOffice转换失败,尝试其他方法")

            except Exception as e:
                print(f"LibreOffice不可用: {e}")

            # 尝试使用unoconv
            try:
                cmd = f"unoconv -f pdf -o {pdf_filepath} {word_filepath}"
                result = os.system(cmd)

                if os.path.exists(pdf_filepath):
                    print(f"PDF生成成功(unoconv): {pdf_filepath}")
                    device_id = os.path.basename(device_dir)
                    pdf_url = f"{MES_REPORT_BASE_URL}/{device_id}/pdf/{pdf_filename}"
                    return {
                        'success': True,
                        'pdf_filepath': pdf_filepath,
                        'pdf_url': pdf_url,
                        'pdf_filename': pdf_filename
                    }
                else:
                    print(f"unoconv转换失败")

            except Exception as e:
                print(f"unoconv不可用: {e}")

            print("PDF转换失败,请安装LibreOffice或unoconv")
            return {
                'success': False,
                'error': 'PDF转换工具不可用'
            }

        except Exception as e:
            print(f"PDF转换失败: {e}")
            return {
                'success': False,
                'error': str(e)
            }

    def generate_report(self, device_id):
        """生成完整报告(Word + PDF)"""
        print(f"开始生成测试报告: {device_id}")

        # 从Parse获取数据
        report_data = self.get_parse_data(device_id)
        if not report_data:
            print("无法获取测试数据")
            return None

        # 生成Word报告
        word_result = self.generate_word_report(report_data)
        if not word_result.get('success'):
            print("Word报告生成失败")
            return None

        # 转换为PDF
        pdf_result = self.convert_to_pdf(word_result['word_filepath'])

        # 返回结果
        result = {
            'device_id': device_id,
            'station_name': report_data.get('station_name', ''),
            'word_url': word_result.get('word_url', ''),
            'pdf_url': pdf_result.get('pdf_url', ''),
            'word_filename': word_result.get('word_filename', ''),
            'pdf_filename': pdf_result.get('pdf_filename', ''),
            'generated_at': datetime.now().isoformat()
        }

        return result


def main():
    """主函数"""
    import argparse

    parser = argparse.ArgumentParser(description='无人机测试报告生成器 - Word模板替换版本')
    parser.add_argument('--device-id', required=True, help='设备ID')
    parser.add_argument('--session-token', help='Parse Session Token')
    parser.add_argument('--pdf', action='store_true', help='同时生成PDF')

    args = parser.parse_args()

    # 创建报告生成器
    generator = UAVTestReportGenerator(args.session_token)

    # 生成报告
    result = generator.generate_report(args.device_id)

    if result:
        print("\n报告生成成功:")
        print(f"  设备ID: {result['device_id']}")
        print(f"  工位: {result['station_name']}")
        print(f"  Word URL: {result['word_url']}")
        if result['pdf_url']:
            print(f"  PDF URL:  {result['pdf_url']}")
        print(f"  生成时间: {result['generated_at']}")

        # 输出JSON格式(用于MES集成)
        print("\nMES集成数据:")
        print(json.dumps(result, ensure_ascii=False, indent=2))
    else:
        print("报告生成失败")
        sys.exit(1)


if __name__ == '__main__':
    main()
